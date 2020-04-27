#!/usr/bin/python3
# @Author: WLP
# @name: docx文档操作.py
# @date 2020-04-27 9:25

#说明文档  https://python-docx.readthedocs.io/en/latest/index.html

from docx import Document
from docx.shared import Inches

document = Document()

#添加标题
document.add_heading('Document Title', 0)

#添加段落
p = document.add_paragraph('A plain paragraph having some ')
#添加加粗内容
p.add_run('bold').bold = True
#添加内容
p.add_run(' and some ')
#添加斜体样式内容
p.add_run('italic.').italic = True

#添加一级标题
document.add_heading('Heading, level 1', level=1)
document.add_paragraph('Intense quote', style='Intense Quote')

document.add_paragraph(
    'first item in unordered list', style='List Bullet'
)
document.add_paragraph(
    'first item in ordered list', style='List Number'
)
#添加图片
document.add_picture('1.png', width=Inches(5))

records = (
    (3, '101', 'Spam'),
    (7, '422', 'Eggs'),
    (4, '631', 'Spam, spam, eggs, and spam')
)

table = document.add_table(rows=1, cols=3)
hdr_cells = table.rows[0].cells
hdr_cells[0].text = 'Qty'
hdr_cells[1].text = 'Id'
hdr_cells[2].text = 'Desc'
for qty, id, desc in records:
    #添加行
    row_cells = table.add_row().cells
    row_cells[0].text = str(qty)
    row_cells[1].text = id
    row_cells[2].text = desc

#添加分页
document.add_page_break()

document.save('demo.docx')

#打开已有文档
document = Document('demo.docx')
senction = document.sections
print(len(senction))

header = senction[0].header
paragrap = header.paragraphs[0]
paragrap.text = 'add paragraph'
document.save('new-file-name.docx')
